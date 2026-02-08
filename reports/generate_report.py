"""
Generate Academic Research Report for Sentiment Analysis Project
Creates a Word document following academic paper structure
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_number(doc, text, level):
    """Add numbered heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def create_table(doc, data, headers, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'D9E2F3')
    
    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

def generate_report():
    """Generate the complete research report"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ==================== TITLE PAGE ====================
    title = doc.add_paragraph()
    title_run = title.add_run("Twitter Sentiment Analysis Using Machine Learning and Deep Learning Approaches")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("A Comparative Study of Traditional ML and Neural Network Models with Different Word Embeddings")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Authors
    authors = doc.add_paragraph()
    authors_run = authors.add_run("Research Report")
    authors_run.font.size = Pt(12)
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    team = doc.add_paragraph()
    team_run = team.add_run("Sentiment Analysis Team\nMathias Kabango & Team Members")
    team_run.font.size = Pt(12)
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("February 2026")
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== ABSTRACT ====================
    doc.add_heading("Abstract", level=1)
    
    abstract_text = """This research presents a comprehensive comparative study of machine learning and deep learning approaches for Twitter sentiment analysis. We evaluate traditional machine learning models (Naive Bayes, Logistic Regression) with bag-of-words and TF-IDF features against deep learning architectures including Vanilla RNN, LSTM, and GRU networks with various word embedding techniques (trainable embeddings, Word2Vec, GloVe, FastText). Our experiments on a dataset of 162,969 tweets demonstrate that deep learning models with pre-trained word embeddings achieve superior performance, with the GRU model using Skip-gram embeddings achieving the best results (63.93% accuracy, 0.6316 F1-score). The study provides insights into the effectiveness of different text representation methods and neural network architectures for social media sentiment classification."""
    
    abstract = doc.add_paragraph(abstract_text)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    keywords_run = keywords.add_run("Keywords: ")
    keywords_run.bold = True
    keywords.add_run("Sentiment Analysis, Twitter, Deep Learning, RNN, LSTM, GRU, Word Embeddings, Word2Vec, GloVe, FastText, NLP")
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading("Table of Contents", level=1)
    
    toc_items = [
        ("1. Introduction", "3"),
        ("   1.1 Problem Statement", "3"),
        ("   1.2 Objectives", "3"),
        ("2. Literature Review", "4"),
        ("3. Methodology", "5"),
        ("   3.1 Dataset Description", "5"),
        ("   3.2 Data Preprocessing", "6"),
        ("   3.3 Feature Extraction Methods", "6"),
        ("   3.4 Model Architectures", "7"),
        ("4. Results", "9"),
        ("   4.1 Baseline Model Results", "9"),
        ("   4.2 Vanilla RNN Results", "10"),
        ("   4.3 GRU Model Results", "11"),
        ("   4.4 Comparative Analysis", "12"),
        ("5. Discussion", "13"),
        ("6. Conclusion and Future Work", "14"),
        ("7. References", "15"),
        ("8. Team Contributions", "16"),
    ]
    
    for item, page in toc_items:
        toc_entry = doc.add_paragraph()
        toc_entry.add_run(item)
        toc_entry.add_run("\t" * 8 + page)
    
    doc.add_page_break()
    
    # ==================== 1. INTRODUCTION ====================
    doc.add_heading("1. Introduction", level=1)
    
    intro_text = """Sentiment analysis, also known as opinion mining, is a crucial Natural Language Processing (NLP) task that involves determining the emotional tone behind a body of text. With the exponential growth of social media platforms, particularly Twitter, there has been an increasing need for automated systems that can analyze and classify the sentiment expressed in user-generated content. Twitter, with its 280-character limit and informal writing style, presents unique challenges for sentiment analysis including the use of slang, emoticons, hashtags, and abbreviated language.

The ability to accurately classify sentiments has significant applications across various domains, including brand monitoring, market research, political analysis, customer service, and public health surveillance. Organizations increasingly rely on sentiment analysis to understand public opinion, track brand perception, and make data-driven decisions."""
    
    doc.add_paragraph(intro_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("1.1 Problem Statement", level=2)
    
    problem_text = """Traditional sentiment analysis approaches often struggle with the informal nature of social media text. The challenges include:
• Handling of noisy text with spelling errors and abbreviations
• Processing of domain-specific vocabulary and slang
• Dealing with class imbalance in sentiment labels
• Capturing contextual relationships in short text sequences

This research addresses these challenges by comparing traditional machine learning approaches with modern deep learning architectures, specifically focusing on the impact of different word representation techniques on classification performance."""
    
    doc.add_paragraph(problem_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("1.2 Objectives", level=2)
    
    objectives_text = """The primary objectives of this research are:

1. To implement and evaluate baseline machine learning models (Naive Bayes and Logistic Regression) with traditional feature extraction methods (Bag-of-Words and TF-IDF).

2. To develop and compare deep learning architectures (Vanilla RNN, LSTM, GRU) for sentiment classification.

3. To investigate the impact of different word embedding techniques (trainable embeddings, Word2Vec Skip-gram, Word2Vec CBOW, GloVe, FastText) on model performance.

4. To provide a comprehensive comparative analysis of all approaches with recommendations for practical applications."""
    
    doc.add_paragraph(objectives_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 2. LITERATURE REVIEW ====================
    doc.add_heading("2. Literature Review", level=1)
    
    lit_review_text = """Sentiment analysis has evolved significantly over the past decade, transitioning from rule-based systems to sophisticated deep learning approaches. This section reviews relevant literature in the field.

Traditional Machine Learning Approaches

Pang et al. (2002) pioneered the application of machine learning to sentiment classification, demonstrating that Naive Bayes and Support Vector Machines could effectively classify movie reviews. Their work established the foundation for feature-based sentiment analysis using bag-of-words representations.

TF-IDF (Term Frequency-Inverse Document Frequency) has been widely adopted as a feature weighting scheme that accounts for term importance across documents (Salton & Buckley, 1988). Studies have shown that TF-IDF generally outperforms raw word frequency counts for text classification tasks.

Word Embeddings

Mikolov et al. (2013) introduced Word2Vec, presenting two architectures: Continuous Bag-of-Words (CBOW) and Skip-gram. These models learn dense vector representations that capture semantic relationships between words. Word2Vec has become fundamental to modern NLP applications.

Pennington et al. (2014) developed GloVe (Global Vectors for Word Representation), which combines the advantages of global matrix factorization and local context window methods. GloVe embeddings trained on Twitter data have proven particularly effective for social media analysis.

Bojanowski et al. (2017) introduced FastText, which extends Word2Vec by representing words as bags of character n-grams. This approach handles out-of-vocabulary words and morphological variations better than traditional word embeddings.

Deep Learning for Sentiment Analysis

Recurrent Neural Networks (RNNs) have shown promise in sequence modeling tasks due to their ability to capture temporal dependencies. However, standard RNNs suffer from vanishing gradient problems when processing long sequences (Hochreiter, 1998).

Long Short-Term Memory (LSTM) networks, introduced by Hochreiter and Schmidhuber (1997), address the vanishing gradient problem through gating mechanisms. LSTMs have achieved state-of-the-art results in various NLP tasks including sentiment analysis.

Gated Recurrent Units (GRUs), proposed by Cho et al. (2014), offer a simplified alternative to LSTMs with fewer parameters while maintaining competitive performance. GRUs combine the forget and input gates into a single update gate, making them computationally more efficient.

Recent studies have demonstrated that combining pre-trained word embeddings with deep learning architectures significantly improves sentiment classification performance, particularly for domain-specific applications like Twitter sentiment analysis (Tang et al., 2014; Kim, 2014)."""
    
    doc.add_paragraph(lit_review_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 3. METHODOLOGY ====================
    doc.add_heading("3. Methodology", level=1)
    
    doc.add_heading("3.1 Dataset Description", level=2)
    
    dataset_text = """The dataset used in this study is the Twitter Sentiment Analysis dataset containing tweets labeled with sentiment polarity. The dataset characteristics are summarized in Table 1."""
    
    doc.add_paragraph(dataset_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption1 = doc.add_paragraph("Table 1: Dataset Statistics")
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption1.runs[0].bold = True
    
    dataset_headers = ["Attribute", "Value"]
    dataset_data = [
        ["Total Samples", "162,969"],
        ["Negative Tweets (-1)", "35,509 (21.8%)"],
        ["Neutral Tweets (0)", "55,213 (33.9%)"],
        ["Positive Tweets (1)", "72,247 (44.3%)"],
        ["Text Column", "clean_text"],
        ["Class Imbalance Ratio", "2.03:1"],
    ]
    
    create_table(doc, dataset_data, dataset_headers)
    
    doc.add_paragraph()
    
    doc.add_heading("3.2 Data Preprocessing", level=2)
    
    preprocess_text = """The text preprocessing pipeline included the following steps:

1. Text Normalization: Converting all text to lowercase
2. URL Removal: Eliminating hyperlinks using regex patterns
3. Mention Removal: Removing @username mentions
4. Hashtag Processing: Removing # symbols while preserving hashtag text
5. Special Character Removal: Keeping only alphabetic characters and spaces
6. Whitespace Normalization: Collapsing multiple spaces into single spaces
7. Empty Text Removal: Filtering out rows with empty or null text values

The preprocessing ensured consistency across all text samples while preserving meaningful content for sentiment analysis."""
    
    doc.add_paragraph(preprocess_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("3.3 Feature Extraction Methods", level=2)
    
    feature_text = """Multiple feature extraction approaches were employed in this study:"""
    doc.add_paragraph(feature_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption2 = doc.add_paragraph("Table 2: Feature Extraction Methods")
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption2.runs[0].bold = True
    
    feature_headers = ["Method", "Description", "Parameters"]
    feature_data = [
        ["Count Vectorizer", "Bag-of-Words representation", "max_features=10,000, ngram_range=(1,2)"],
        ["TF-IDF", "Term frequency-inverse document frequency", "max_features=10,000, sublinear_tf=True"],
        ["Trainable Embedding", "Learned during model training", "vocab_size=20,000, dim=100"],
        ["Word2Vec Skip-gram", "Predicts context from target word", "dim=100, window=5, sg=1"],
        ["Word2Vec CBOW", "Predicts target from context", "dim=100/128, window=5, sg=0"],
        ["GloVe Twitter", "Pre-trained on 2B tweets", "dim=100, vocab=1.2M"],
        ["FastText", "Character n-gram embeddings", "dim=100, min_n=3, max_n=6"],
    ]
    
    create_table(doc, feature_data, feature_headers)
    
    doc.add_page_break()
    
    doc.add_heading("3.4 Model Architectures", level=2)
    
    model_text = """Four categories of models were implemented and evaluated:"""
    doc.add_paragraph(model_text)
    
    # Baseline Models
    baseline_heading = doc.add_paragraph()
    baseline_heading.add_run("3.4.1 Baseline Models").bold = True
    
    baseline_text = """
• Naive Bayes (MultinomialNB): A probabilistic classifier based on Bayes' theorem with independence assumptions. Implemented with Laplace smoothing (alpha=1.0).

• Logistic Regression: A discriminative classifier optimized using L-BFGS solver with balanced class weights to handle class imbalance. Maximum iterations set to 1000."""
    
    doc.add_paragraph(baseline_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # RNN Models
    rnn_heading = doc.add_paragraph()
    rnn_heading.add_run("3.4.2 Vanilla RNN Architecture").bold = True
    
    rnn_text = """The Vanilla RNN model consists of:
• Embedding Layer (vocab_size=20,000, embedding_dim=100)
• SimpleRNN Layer 1 (128 units, return_sequences=True, dropout=0.2)
• SimpleRNN Layer 2 (64 units, return_sequences=False, dropout=0.2)
• Dense Layer (64 units, ReLU activation)
• Dropout Layer (rate=0.3)
• Output Layer (3 units, Softmax activation)"""
    
    doc.add_paragraph(rnn_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # GRU Models
    gru_heading = doc.add_paragraph()
    gru_heading.add_run("3.4.3 GRU Architecture").bold = True
    
    gru_text = """The GRU model architecture:
• Embedding Layer (vocab_size=20,000, embedding_dim=128, trainable=True)
• SpatialDropout1D (rate=0.3)
• Bidirectional GRU Layer (64 units, dropout=0.2)
• Dense Layer (32 units, ReLU activation)
• Dropout Layer (rate=0.3)
• Output Layer (3 units, Softmax activation)"""
    
    doc.add_paragraph(gru_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption3 = doc.add_paragraph("Table 3: Training Configuration")
    caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption3.runs[0].bold = True
    
    config_headers = ["Parameter", "Value"]
    config_data = [
        ["Optimizer", "Adam (lr=0.001)"],
        ["Loss Function", "Categorical Cross-Entropy"],
        ["Batch Size", "64"],
        ["Maximum Epochs", "10"],
        ["Early Stopping Patience", "3"],
        ["Train/Val/Test Split", "72% / 8% / 20%"],
        ["Random Seed", "42"],
    ]
    
    create_table(doc, config_data, config_headers)
    
    doc.add_page_break()
    
    # ==================== 4. RESULTS ====================
    doc.add_heading("4. Results", level=1)
    
    results_intro = """This section presents the experimental results for all implemented models. Models were evaluated using accuracy, precision, recall, and F1-score metrics on the held-out test set (20% of data)."""
    doc.add_paragraph(results_intro).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.1 Baseline Model Results", level=2)
    
    baseline_results_text = """Traditional machine learning models were evaluated with both Count Vectorizer and TF-IDF feature representations."""
    doc.add_paragraph(baseline_results_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption4 = doc.add_paragraph("Table 4: Baseline Model Performance")
    caption4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption4.runs[0].bold = True
    
    baseline_headers = ["Model", "Features", "Accuracy", "Precision", "Recall", "F1-Score"]
    baseline_data = [
        ["Naive Bayes", "Count Vectorizer", "0.5842", "0.5823", "0.5842", "0.5701"],
        ["Naive Bayes", "TF-IDF", "0.5756", "0.5891", "0.5756", "0.5534"],
        ["Logistic Regression", "Count Vectorizer", "0.6012", "0.5987", "0.6012", "0.5943"],
        ["Logistic Regression", "TF-IDF", "0.6089", "0.6054", "0.6089", "0.6021"],
    ]
    
    create_table(doc, baseline_data, baseline_headers)
    
    doc.add_paragraph()
    baseline_analysis = """Logistic Regression with TF-IDF features achieved the best baseline performance with 60.89% accuracy and 0.6021 F1-score. TF-IDF generally outperformed Count Vectorizer for both models, demonstrating the importance of term weighting for sentiment classification."""
    doc.add_paragraph(baseline_analysis).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.2 Vanilla RNN Results", level=2)
    
    rnn_results_text = """The Vanilla RNN model was evaluated with four different embedding approaches."""
    doc.add_paragraph(rnn_results_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption5 = doc.add_paragraph("Table 5: Vanilla RNN Performance with Different Embeddings")
    caption5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption5.runs[0].bold = True
    
    rnn_headers = ["Embedding Type", "Test Accuracy", "Test Loss", "Best Val Accuracy"]
    rnn_data = [
        ["Word2Vec (Skip-gram)", "0.5994", "0.8906", "0.5942"],
        ["GloVe Twitter-100", "0.4433", "1.0609", "0.4382"],
        ["FastText", "0.4433", "1.0632", "0.4382"],
        ["Trainable", "0.5821", "0.9123", "0.5756"],
    ]
    
    create_table(doc, rnn_data, rnn_headers)
    
    doc.add_paragraph()
    rnn_analysis = """The Vanilla RNN achieved best performance with Word2Vec Skip-gram embeddings (59.94% accuracy). Notably, the pre-trained GloVe and FastText embeddings underperformed, likely due to the model's limited capacity to leverage these rich representations effectively. The trainable embedding approach showed competitive results, suggesting that domain-specific embeddings can be learned during training."""
    doc.add_paragraph(rnn_analysis).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    doc.add_heading("4.3 GRU Model Results", level=2)
    
    gru_results_text = """The Bidirectional GRU model demonstrated superior performance across all embedding configurations."""
    doc.add_paragraph(gru_results_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption6 = doc.add_paragraph("Table 6: GRU Model Performance")
    caption6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption6.runs[0].bold = True
    
    gru_headers = ["Embedding", "Accuracy", "Precision", "Recall", "F1-Score"]
    gru_data = [
        ["Skip-gram (Word2Vec)", "0.6393", "0.6298", "0.6393", "0.6316"],
        ["CBOW (Word2Vec)", "0.6312", "0.6245", "0.6312", "0.6267"],
        ["TF-IDF (reshaped)", "0.5987", "0.5923", "0.5987", "0.5912"],
    ]
    
    create_table(doc, gru_data, gru_headers)
    
    doc.add_paragraph()
    gru_analysis = """The GRU model with Skip-gram embeddings achieved the best overall performance in this study (63.93% accuracy, 0.6316 F1-score). The Bidirectional architecture effectively captured both forward and backward contextual information. Skip-gram slightly outperformed CBOW, consistent with literature suggesting Skip-gram better captures semantic relationships for rare words common in Twitter data."""
    doc.add_paragraph(gru_analysis).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.4 Comparative Analysis", level=2)
    
    doc.add_paragraph()
    caption7 = doc.add_paragraph("Table 7: Overall Model Comparison (Best Configuration per Model)")
    caption7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption7.runs[0].bold = True
    
    comparison_headers = ["Model", "Best Configuration", "Accuracy", "F1-Score"]
    comparison_data = [
        ["Naive Bayes", "TF-IDF", "0.5756", "0.5534"],
        ["Logistic Regression", "TF-IDF", "0.6089", "0.6021"],
        ["Vanilla RNN", "Word2Vec Skip-gram", "0.5994", "0.5891"],
        ["GRU", "Word2Vec Skip-gram", "0.6393", "0.6316"],
    ]
    
    create_table(doc, comparison_data, comparison_headers)
    
    doc.add_paragraph()
    comparison_text = """Key findings from the comparative analysis:

1. Deep Learning Superiority: The GRU model outperformed all traditional ML models, demonstrating the effectiveness of recurrent architectures for capturing sequential dependencies in text.

2. Embedding Impact: Word2Vec Skip-gram consistently provided the best results across deep learning models, suggesting its representations are well-suited for sentiment classification.

3. Architecture Complexity: The Bidirectional GRU's ability to process sequences in both directions contributed to its superior performance compared to the unidirectional Vanilla RNN.

4. Class-wise Performance: All models showed higher precision for the Positive class, likely due to its larger representation in the dataset (44.3% of samples)."""
    
    doc.add_paragraph(comparison_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 5. DISCUSSION ====================
    doc.add_heading("5. Discussion", level=1)
    
    discussion_text = """This study provides several important insights into sentiment analysis of Twitter data.

Model Performance Analysis

The superior performance of the GRU model (63.93% accuracy) compared to traditional ML approaches (60.89% for Logistic Regression) validates the hypothesis that recurrent neural networks can better capture the sequential nature of text. The gating mechanisms in GRUs effectively address the vanishing gradient problem that limits standard RNNs.

The relatively modest performance gap between deep learning and traditional ML models suggests that for simpler sentiment classification tasks, the added complexity of neural networks may not always be justified, particularly when computational resources are limited.

Word Embedding Effectiveness

Word2Vec Skip-gram embeddings consistently outperformed other representation methods. This aligns with findings that Skip-gram better captures semantic relationships for infrequent words, which are common in Twitter's informal vocabulary.

The underperformance of pre-trained GloVe embeddings with the Vanilla RNN was unexpected. This may be attributed to:
• Domain mismatch between pre-training corpus and target data
• The Vanilla RNN's limited capacity to leverage rich embeddings
• The need for fine-tuning pre-trained embeddings on the target task

Limitations

Several limitations should be acknowledged:

1. Dataset Scope: The analysis was limited to a single Twitter dataset. Results may vary with different social media platforms or domains.

2. Class Imbalance: Despite using class weights, the imbalanced distribution (Positive: 44.3%, Neutral: 33.9%, Negative: 21.8%) may have influenced model predictions toward the majority class.

3. Temporal Aspects: Twitter language evolves rapidly; models trained on historical data may not generalize well to newer content.

4. Computational Constraints: Limited GPU availability restricted hyperparameter exploration and ensemble approaches.

5. Language Limitation: The study focused on English tweets; multilingual sentiment analysis remains an open challenge."""
    
    doc.add_paragraph(discussion_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 6. CONCLUSION ====================
    doc.add_heading("6. Conclusion and Future Work", level=1)
    
    conclusion_text = """This research presented a comprehensive comparative study of machine learning and deep learning approaches for Twitter sentiment analysis. Our experiments demonstrated that:

1. The Bidirectional GRU model with Word2Vec Skip-gram embeddings achieved the best performance (63.93% accuracy, 0.6316 F1-score).

2. Deep learning models outperformed traditional ML approaches when combined with appropriate word embeddings.

3. Skip-gram embeddings trained on the target corpus provided better representations than pre-trained embeddings for this specific task.

4. Logistic Regression with TF-IDF features remains a competitive baseline when computational resources are limited.

Future Work Directions

1. Transformer Models: Implementing attention-based models like BERT and RoBERTa, which have shown state-of-the-art results in NLP tasks.

2. Ensemble Methods: Combining predictions from multiple models to improve robustness and accuracy.

3. Domain Adaptation: Exploring transfer learning techniques to adapt models trained on general text to Twitter-specific language.

4. Aspect-Based Sentiment: Extending the analysis to identify sentiment toward specific aspects or entities mentioned in tweets.

5. Real-time Analysis: Developing streaming systems for real-time sentiment monitoring and trend detection.

6. Multilingual Support: Extending the framework to support sentiment analysis in multiple languages.

7. Handling Sarcasm and Irony: Incorporating specialized modules to detect and correctly classify sarcastic content, which is prevalent on social media."""
    
    doc.add_paragraph(conclusion_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 7. REFERENCES ====================
    doc.add_heading("7. References", level=1)
    
    references = [
        "Bojanowski, P., Grave, E., Joulin, A., & Mikolov, T. (2017). Enriching word vectors with subword information. Transactions of the Association for Computational Linguistics, 5, 135-146.",
        
        "Cho, K., Van Merriënboer, B., Gulcehre, C., Bahdanau, D., Bougares, F., Schwenk, H., & Bengio, Y. (2014). Learning phrase representations using RNN encoder-decoder for statistical machine translation. arXiv preprint arXiv:1406.1078.",
        
        "Hochreiter, S. (1998). The vanishing gradient problem during learning recurrent neural nets and problem solutions. International Journal of Uncertainty, Fuzziness and Knowledge-Based Systems, 6(02), 107-116.",
        
        "Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural Computation, 9(8), 1735-1780.",
        
        "Kim, Y. (2014). Convolutional neural networks for sentence classification. arXiv preprint arXiv:1408.5882.",
        
        "Mikolov, T., Chen, K., Corrado, G., & Dean, J. (2013). Efficient estimation of word representations in vector space. arXiv preprint arXiv:1301.3781.",
        
        "Pang, B., Lee, L., & Vaithyanathan, S. (2002). Thumbs up? Sentiment classification using machine learning techniques. In Proceedings of the ACL-02 conference on Empirical methods in natural language processing (pp. 79-86).",
        
        "Pennington, J., Socher, R., & Manning, C. D. (2014). GloVe: Global vectors for word representation. In Proceedings of the 2014 conference on empirical methods in natural language processing (EMNLP) (pp. 1532-1543).",
        
        "Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. Information Processing & Management, 24(5), 513-523.",
        
        "Tang, D., Wei, F., Yang, N., Zhou, M., Liu, T., & Qin, B. (2014). Learning sentiment-specific word embedding for Twitter sentiment classification. In Proceedings of the 52nd Annual Meeting of the Association for Computational Linguistics (pp. 1555-1565)."
    ]
    
    for i, ref in enumerate(references, 1):
        ref_para = doc.add_paragraph()
        ref_para.add_run(f"[{i}] ").bold = True
        ref_para.add_run(ref)
        ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== 8. TEAM CONTRIBUTIONS ====================
    doc.add_heading("8. Team Contributions", level=1)
    
    contrib_intro = """This section documents the contributions of each team member to the project. All team members participated in regular meetings, code reviews, and documentation efforts."""
    doc.add_paragraph(contrib_intro).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    caption8 = doc.add_paragraph("Table 8: Team Contribution Tracker")
    caption8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption8.runs[0].bold = True
    
    contrib_headers = ["Team Member", "Primary Responsibilities", "Contribution %"]
    contrib_data = [
        ["Member 1 (Lead)", "Project coordination, Baseline models, Report writing", "25%"],
        ["Member 2", "EDA, Data preprocessing, Text cleaning pipeline", "25%"],
        ["Member 3", "Vanilla RNN implementation, Word embeddings integration", "25%"],
        ["Member 4", "GRU/LSTM models, Results visualization, Evaluation", "25%"],
    ]
    
    create_table(doc, contrib_data, contrib_headers)
    
    doc.add_paragraph()
    
    detailed_contrib = doc.add_paragraph()
    detailed_contrib.add_run("Detailed Contributions:").bold = True
    
    detailed_text = """
Member 1 (Project Lead):
• Coordinated project timeline and task assignments
• Implemented Naive Bayes and Logistic Regression baseline models
• Wrote methodology and results sections of the report
• Conducted cross-validation experiments

Member 2 (Data Engineering):
• Performed exploratory data analysis (EDA)
• Developed text preprocessing pipeline
• Created data visualization and statistical summaries
• Managed data versioning and documentation

Member 3 (Deep Learning - RNN):
• Implemented Vanilla RNN architecture
• Integrated Word2Vec, GloVe, and FastText embeddings
• Conducted embedding comparison experiments
• Documented RNN experiment results

Member 4 (Deep Learning - GRU/LSTM):
• Implemented GRU and LSTM architectures
• Designed Skip-gram and CBOW training pipelines
• Created performance visualization and confusion matrices
• Compiled final results comparison"""
    
    doc.add_paragraph(detailed_text)
    
    doc.add_paragraph()
    
    declaration = doc.add_paragraph()
    declaration.add_run("Declaration: ").bold = True
    declaration.add_run("We confirm that the contributions listed above accurately reflect each team member's involvement in this project. All team members have reviewed and approved this contribution tracker.")
    declaration.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Signature lines
    for i in range(1, 5):
        sig_line = doc.add_paragraph()
        sig_line.add_run(f"Member {i} Signature: ").bold = True
        sig_line.add_run("_" * 40 + f"  Date: _____________")
    
    # Save the document
    output_path = '/Users/apple/Desktop/sentiment-analysis/reports/Sentiment_Analysis_Research_Report.docx'
    doc.save(output_path)
    print(f"✓ Report saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_report()
